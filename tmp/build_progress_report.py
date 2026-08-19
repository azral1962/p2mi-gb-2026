from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "Laporan_Kemajuan_Riset_TISE_3.0_2026.docx"
PDF_PATH = OUT_DIR / "Laporan_Kemajuan_Riset_TISE_3.0_2026.pdf"

NAVY = "1F4D78"
BLUE = "2E74B5"
MUTED = "5E6B75"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF7E0"
WHITE = "FFFFFF"
BLACK = "111111"
RED = "9B1C1C"
GREEN = "2F6B4F"


PUBLICATIONS = [
    (1, "Correction: Orchestrating value co-creation in digital education using the TISE-VALORIZE framework: from knowing to becoming", "Frontiers in Education", "Q1", "Published"),
    (2, "Orchestrating value co-creation in digital education using the TISE-VALORIZE framework: from knowing to becoming", "Frontiers in Education", "Q1", "Published"),
    (3, "A Hybrid Systematic Literature Review of User-Centric QoE and QoS Evaluation Methods in Video Streaming Services", "JOIV: International Journal on Informatics Visualization", "Q2/Q3", "Published"),
    (4, "A Synergy-Driven Portfolio Optimization Model Under Mixed Financing", "JOIV: International Journal on Informatics Visualization", "Q2/Q3", "Published"),
    (5, "An Integrated Practical Assessment Model for Web Programming Education in Smart Campus Environments", "Ingenierie des Systemes d'Information", "Q3", "Published"),
    (6, "A Smart Management Framework for Higher Education Accreditation Based on the PUDAL Engine and V-Model Approach", "Asia-Pacific Journal of Information Technology & Multimedia", "Q3", "Published"),
    (7, "Focus Detection Using EEG: Trends, Challenges, Advantages, Application Areas, Background", "Health Information Science: 14th International Conference, HIS 2025", "Non-Q proceedings", "Published"),
    (8, "Structured Collaborative Problem Solving Using Isomorphic Task Design: Evidence on Engineering Student Outcomes", "IEEE Transactions on Education", "Q1", "Under review"),
    (9, "GRACE and MyGRACE: Designing and Evaluating a Narrative-Centered Socio-Technical Ecosystem for Older-Adult Wellbeing", "IEEE Transactions on Technology and Society", "Q1", "Draft"),
    (10, "Managing Pre-Deployment Mission-Oriented Service Ecosystems: A Multi-Stakeholder Recommendation Core and Digital-Twin Design Method", "IEEE Transactions on Engineering Management", "Q1", "Draft"),
]

PROGRESS_ROWS = [
    (
        "Landasan teori dan meta-arsitektur",
        "Formulasi TISE 3.0 sebagai meta-artefak berbasis MAS.",
        "Manuskrip SVCC memuat prompt-to-artifact, ASTF, PSKVE/VOMR, tiga peran agen, dan W-Model. [S3]",
        "Maju substansial",
    ),
    (
        "Purwarupa dan integrasi",
        "Konversi prompt berentropi tinggi menuju spesifikasi SysML v2 dan artefak TISE 2.0.",
        "SVCC mendeskripsikan paket Python modular, reasoner, coordinator, tiga pasar, provider LLM, serta alur eksekusi Mini-Lab. Kode tidak menjadi lampiran yang diperiksa dalam laporan ini. [S3]",
        "Dalam proses",
    ),
    (
        "Validasi PICOC dan Digital Twin",
        "Skenario, komparator, metrik, dan pengujian kompresi nilai serta resiliensi.",
        "Desain eksperimen, pertanyaan validasi, skenario, dan artefak hasil sudah disiapkan. Bagian Results masih menyatakan data simulasi belum tersedia. [S3]",
        "Siap validasi",
    ),
    (
        "Luaran publikasi Q1",
        "Satu artikel Q1 dengan target submit 30 November 2026. [S1]",
        "Pipeline mencatat 7 published, 1 under review, dan 2 draft. Manuskrip SVCC sangat relevan langsung, tetapi belum tercatat dengan target jurnal/status submit. [S2][S3]",
        "Belum terpenuhi",
    ),
]

FINANCE_ROWS = [
    ("Pegawai", 59_040_000, 26_240_000, "44,4%", "Empat pembayaran honor kegiatan April-Juli."),
    ("Barang", 23_340_000, 15_840_000, "67,9%", "Dua transaksi aksesori komputer."),
    ("Jasa", 67_620_000, 29_000_000, "42,9%", "Honor jasa civitas, sewa alat, dan pengolahan data."),
    ("Modal", 0, 10_400_000, "n/a", "Belanja modal belum tampil sebagai pos terpisah dalam RAB proposal."),
]

WORKPLAN_ROWS = [
    ("Agustus-September", "Kunci PICOC, skenario, seed, ambang penerimaan, dan traceability matrix; audit kesesuaian implementasi dengan arsitektur."),
    ("September-Oktober", "Jalankan eksperimen berpasangan baseline-SOTA-TISE; lakukan replikasi, interval ketidakpastian, sensitivitas, ablation, dan uji gangguan."),
    ("Oktober-November", "Isi Results dan Discussion dengan hasil final; tetapkan jurnal sasaran; lengkapi acknowledgement PPMI; submit paling lambat 30 November 2026."),
    ("Desember", "Tindak lanjut editorial, arsipkan artefak replikasi, selesaikan laporan akhir, dan rekonsiliasi klasifikasi serta bukti penggunaan dana."),
]


def rupiah(value: int) -> str:
    return "Rp" + f"{value:,}".replace(",", ".")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="B7C0C8", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_in: Sequence[float], indent_dxa: int = 120) -> None:
    widths_dxa = [round(v * 1440) for v in widths_in]
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=11, bold=False, color=BLACK, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])
    set_run_font(run, size=9, color=MUTED)


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "LAPORAN KEMAJUAN RISET PPMI - RISET GURU BESAR 2026"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.runs[0], size=8.5, bold=True, color=MUTED)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        r = fp.add_run("Halaman ")
        set_run_font(r, size=9, color=MUTED)
        add_field(fp, "PAGE")


def add_docx_para(doc: Document, text: str, bold_lead: str | None = None, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)
    return p


def add_docx_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[float], font_size=9.0, header_fill=LIGHT):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if idx == 0 and len(headers) > 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def build_docx() -> None:
    doc = Document()
    set_doc_styles(doc)
    add_header_footer(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("LAPORAN KEMAJUAN PROYEK RISET")
    set_run_font(r, size=23, bold=True, color=BLACK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Meta-Architecture of Triune-Intelligence Smart Engineering (TISE 3.0)")
    set_run_font(r, size=14, bold=True, color=NAVY)

    for label, value in (
        ("Skema", "Program Riset P2MI STEI - Riset Guru Besar Tahun Anggaran 2026"),
        ("Kode file", "STEI.PPMI-1-43-2026"),
        ("Ketua peneliti", "Prof. Ir. Armein Z. R. Langi, M.Sc., Ph.D."),
        ("Unit", "Kelompok Keahlian Teknologi Informasi, STEI ITB"),
        ("Periode", "1 April - 31 Desember 2026"),
        ("Posisi laporan", "19 Agustus 2026; transaksi keuangan tersedia sampai 28 Juli 2026"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + ": ")
        set_run_font(r, bold=True)
        r = p.add_run(value)
        set_run_font(r)

    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [6.5])
    set_table_borders(callout, color="D7B75D", size="8")
    set_repeat_table_header(callout.rows[0])
    set_cell_shading(callout.cell(0, 0), PALE_GOLD)
    cp = callout.cell(0, 0).paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    cr = cp.add_run("STATUS KESELURUHAN: BERJALAN SESUAI TAHAPAN, DENGAN RISIKO TERKENDALI")
    set_run_font(cr, size=11, bold=True, color=NAVY)
    cp = callout.cell(0, 0).add_paragraph()
    cp.paragraph_format.space_after = Pt(0)
    cr = cp.add_run("Arsitektur, rancangan purwarupa, dan instrumen eksperimen telah terdokumentasi. Prioritas berikutnya adalah menghasilkan bukti simulasi final, mengisi bagian hasil, memastikan keterlacakan artefak, dan menuntaskan submit artikel Q1.")
    set_run_font(cr, size=10.5)

    doc.add_heading("1. Ringkasan Eksekutif", level=1)
    add_docx_para(doc, "Riset TISE 3.0 bertujuan membangun meta-arsitektur berbasis prompt yang mengubah niat pemangku kepentingan menjadi spesifikasi SysML v2 dan artefak TISE 2.0 melalui orkestrasi Meta-Architect, Generative Engineer, dan Critic. Berdasarkan dokumen yang tersedia, pekerjaan telah bergerak dari formulasi konsep menuju kesiapan validasi Digital Twin. [S1][S3]")
    metrics = [
        ("Kemajuan teknis", "Arsitektur dan desain eksperimen tersedia; hasil empiris final belum tersedia."),
        ("Luaran ilmiah", "Pipeline: 7 published, 1 under review, 2 draft; terdapat pula manuskrip SVCC 74 halaman yang relevan langsung."),
        ("Keuangan", "Realisasi Rp81.480.000 dari Rp150.000.000; serapan 54% menurut LPD (54,32% dari perhitungan)."),
        ("Target kontraktual", "Satu publikasi Q1; rencana submit 30 November 2026."),
    ]
    add_docx_table(doc, ["Aspek", "Posisi Kemajuan"], metrics, [1.55, 4.95], font_size=9.5, header_fill=PALE_BLUE)

    doc.add_heading("2. Sasaran dan Posisi Terhadap Jadwal", level=1)
    add_docx_para(doc, "Proposal menetapkan tiga tujuan: memformulasikan TISE 3.0, membangun dan mensimulasikan purwarupa prompt-to-artifact, serta mempublikasikan hasil pada jurnal internasional Q1. Jadwal April-Desember menempatkan Agustus-September pada tahap Inner Right W-Model, yaitu persiapan validasi PICOC dan Digital Twin. [S1]")
    add_docx_table(
        doc,
        ["Paket Kerja", "Target Proposal", "Bukti Kemajuan", "Status"],
        PROGRESS_ROWS,
        [1.25, 1.55, 2.75, 0.95],
        font_size=8.4,
    )
    add_docx_para(doc, "Catatan penilaian: status di atas diturunkan dari dokumen, bukan dari pemeriksaan repositori kode atau log eksperimen. Klaim implementasi yang disebut dalam SVCC perlu dilengkapi dengan versi kode, konfigurasi, seed, dan keluaran run yang dibekukan.", italic=True)

    doc.add_heading("3. Capaian Teknis dan Ilmiah", level=1)
    doc.add_heading("3.1 Meta-arsitektur TISE 3.0", level=2)
    add_docx_para(doc, "Manuskrip Systemic Value Co-Creation (SVCC) mendefinisikan TISE 3.0 sebagai lapisan design-time yang menghasilkan artefak sosio-teknis TISE 2.0. Transformasi disusun dalam lima tahap: prompt pemangku kepentingan, dekomposisi ASTF, formalisasi PSKVE/VOMR, spesifikasi SysML v2, dan instansiasi artefak. Peran agen dipisahkan agar keluaran antara, konflik, dan pemeriksaan dapat diaudit. [S3]")
    doc.add_heading("3.2 Purwarupa dan Digital Twin Mini-Lab", level=2)
    add_docx_para(doc, "SVCC mendeskripsikan paket Python berorientasi objek dengan Stakeholder, VOMR account, beberapa MIG reasoner, tiga coordinator, pasar Stock/Financial/Resource-Output, PSKVE Exchange Engine, Triune Intelligence, dan adapter penyedia LLM. Alur simulasi mencakup observasi, proposal MIG, koordinasi, pemeriksaan kendala, eksekusi, settlement, dan pencatatan EVI/resiliensi. Ini merupakan bukti kesiapan arsitektural dan implementatif, tetapi laporan ini belum memverifikasi berkas kode secara terpisah. [S3]")
    doc.add_heading("3.3 Validasi dan batas klaim", level=2)
    add_docx_para(doc, "Desain eksperimen membedakan efek penalaran stakeholder dari efek koordinasi dan membandingkan baseline, fixed-policy SOTA, serta adaptive TISE. Pertanyaan validasi, skenario normal/gangguan, metrik, dan format tabel hasil telah tersedia. Namun, bagian Results masih berupa kerangka yang akan diisi setelah Mini-Lab dijalankan; karena itu belum layak menyatakan peningkatan performa, resiliensi, atau efisiensi kompresi nilai sebagai temuan empiris. [S3]")

    doc.add_heading("4. Luaran Publikasi", level=1)
    add_docx_para(doc, "Dokumen Publication Pipeline 2026 mencatat sepuluh luaran: tujuh published, satu under review, dan dua draft. Dari tujuh yang published, dua ditandai Q1, dua Q2/Q3, dua Q3, dan satu prosiding non-Q. [S2]")
    add_docx_table(
        doc,
        ["Status", "Jumlah", "Rincian"],
        [
            ("Published", "7", "2 Q1; 2 Q2/Q3; 2 Q3; 1 prosiding non-Q"),
            ("Under review", "1", "Q1 - IEEE Transactions on Education"),
            ("Draft", "2", "Keduanya diarahkan ke jurnal Q1 IEEE"),
        ],
        [1.35, 0.75, 4.4],
        font_size=9.5,
        header_fill=PALE_BLUE,
    )
    add_docx_para(doc, "Kesesuaian dengan target proposal: target kontraktual adalah satu artikel bertopik TISE 3.0 untuk jurnal Q1 dengan rencana submit 30 November 2026. Tidak ada entri pipeline yang sama persis dengan judul target. Manuskrip SVCC memiliki keterkaitan langsung dan isi yang memadai sebagai basis artikel, tetapi status jurnal/submit belum dinyatakan dan bagian hasil masih menunggu data. Oleh sebab itu target publikasi proyek dinilai masih dalam proses, meskipun pipeline publikasi peneliti secara umum kuat. [S1][S2][S3]")

    doc.add_page_break()
    doc.add_heading("4.1 Daftar Publication Pipeline 2026", level=2)
    pub_rows = [(str(no), title, journal, quartile, status) for no, title, journal, quartile, status in PUBLICATIONS]
    add_docx_table(doc, ["No", "Judul", "Target jurnal/forum", "Kuartil", "Status"], pub_rows, [0.35, 3.05, 1.45, 0.65, 1.0], font_size=8.0)
    add_docx_para(doc, "Catatan: daftar ini direproduksi dari pipeline yang disediakan. Laporan tidak mengasumsikan bahwa seluruh entri merupakan luaran langsung pendanaan PPMI TISE 3.0 tanpa bukti acknowledgement atau pemetaan biaya. [S2]", italic=True)

    doc.add_heading("5. Penggunaan Dana", level=1)
    add_docx_para(doc, "LPD mencatat penerimaan Rp150.000.000 dan 13 transaksi pengeluaran dengan total Rp81.480.000. Saldo Rp68.520.000 dan presentase serapan dilaporkan 54%. Perhitungan terhadap pagu menghasilkan 54,32%, sehingga perbedaan hanya akibat pembulatan. [S4]")
    add_docx_table(
        doc,
        ["Indikator", "Nilai"],
        [
            ("Pagu/penerimaan", rupiah(150_000_000)),
            ("Realisasi", rupiah(81_480_000)),
            ("Saldo", rupiah(68_520_000)),
            ("Serapan", "54% (LPD); 54,32% (perhitungan)"),
            ("Cakupan transaksi", "28 April - 28 Juli 2026; laporan ditandatangani/dinyatakan Agustus 2026"),
        ],
        [2.0, 4.5],
        font_size=9.5,
        header_fill=PALE_BLUE,
    )

    finance_table = []
    for category, planned, actual, absorption, note in FINANCE_ROWS:
        finance_table.append((category, rupiah(planned) if planned else "Tidak terpisah", rupiah(actual), absorption, note))
    finance_table.append(("Total", rupiah(150_000_000), rupiah(81_480_000), "54,32%", "Saldo keseluruhan " + rupiah(68_520_000) + "."))
    add_docx_table(doc, ["Kategori", "RAB Proposal", "Realisasi LPD", "Serapan", "Keterangan"], finance_table, [1.05, 1.2, 1.2, 0.75, 2.3], font_size=8.3)
    add_docx_para(doc, "Catatan rekonsiliasi: perbandingan kategori bersifat indikatif. RAB proposal menggabungkan honor mahasiswa, perjalanan, AI/hosting, internet/cloud, serta registrasi/publikasi di kelompok jasa; LPD menggunakan klasifikasi pegawai, barang, jasa, dan modal. Belanja modal Rp10.400.000 perlu dipetakan secara administratif ke pos RAB yang disetujui sebelum laporan akhir. [S1][S4]", italic=True)

    doc.add_heading("6. Isu, Risiko, dan Tindakan Korektif", level=1)
    risk_rows = [
        ("Bukti empiris belum final", "Tinggi", "Bekukan protokol, jalankan repeated matched runs, laporkan interval ketidakpastian, sensitivitas, ablation, dan kasus gagal."),
        ("Target Q1 belum terdokumentasi sebagai submitted", "Tinggi", "Tetapkan artikel utama dan jurnal, selesaikan Results/Discussion, pemeriksaan format dan acknowledgement, submit sebelum 30 November."),
        ("Keterlacakan prompt-SysML-artefak", "Sedang", "Sediakan trace matrix, versi model/kode, konfigurasi provider, seed, log, dan artefak keluaran yang dapat direplikasi."),
        ("Klasifikasi belanja modal", "Sedang", "Rekonsiliasi dengan RAB dan bagian keuangan; dokumentasikan justifikasi serta bukti persetujuan bila diperlukan."),
        ("Batas klaim simulasi", "Sedang", "Pisahkan bukti simulasi dari klaim kondisi kampus nyata; jelaskan asumsi, kalibrasi, dan kebutuhan validasi eksternal."),
    ]
    add_docx_table(doc, ["Isu/Risiko", "Prioritas", "Tindakan"], risk_rows, [1.65, 0.8, 4.05], font_size=8.7)

    doc.add_heading("7. Rencana Kerja Sampai Akhir Tahun", level=1)
    add_docx_table(doc, ["Periode", "Keluaran yang Diharapkan"], WORKPLAN_ROWS, [1.45, 5.05], font_size=9.2, header_fill=PALE_BLUE)

    doc.add_heading("8. Kesimpulan", level=1)
    add_docx_para(doc, "Secara keseluruhan, riset berada pada posisi yang konsisten dengan peralihan menuju tahap validasi Agustus-September. Fondasi konseptual, rancangan meta-arsitektur, struktur purwarupa, dan protokol Digital Twin telah terdokumentasi dengan baik. Kesenjangan utama adalah belum tersedianya hasil eksperimen final dan belum adanya bukti bahwa artikel TISE 3.0 telah masuk proses submit. Dengan serapan anggaran 54% dan sisa periode sampai Desember, sasaran masih dapat dicapai apabila eksekusi eksperimen, penyusunan artikel, serta rekonsiliasi anggaran diprioritaskan segera.")

    doc.add_page_break()
    doc.add_heading("Lampiran: Sumber Data dan Catatan Verifikasi", level=1)
    sources = [
        ("[S1]", "proposal5signed.pdf - proposal 18 halaman; identitas, tujuan, target publikasi, jadwal, dan RAB."),
        ("[S2]", "bahan/pipeline2026.pdf - Publication Pipeline 2026, 2 halaman."),
        ("[S3]", "bahan/SVCC.pdf - Systemic Value Co-Creation, 74 halaman, bertanggal 19 Agustus 2026."),
        ("[S4]", "LPD Armein Z.R Langi.pdf - Laporan Penggunaan Dana, 2 halaman, periode 1 April-31 Desember 2026."),
    ]
    add_docx_table(doc, ["Kode", "Dokumen"], sources, [0.65, 5.85], font_size=9.0)
    add_docx_para(doc, "Metode: angka dan status ditranskripsi dari dokumen yang diberikan dan diperiksa terhadap tampilan halaman PDF. Tidak dilakukan verifikasi eksternal atas indeks jurnal, tautan bukti transfer, atau keberadaan kode sumber. Istilah 'published', 'under review', dan 'draft' mengikuti pipeline. Status SVCC dinilai sebagai manuskrip pra-submisi karena tidak mencantumkan jurnal/status submit dan bagian Results menyatakan data belum tersedia.", italic=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Bandung, 19 Agustus 2026")
    set_run_font(r)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(40)
    r = p.add_run("Ketua Peneliti,")
    set_run_font(r)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Prof. Ir. Armein Z. R. Langi, M.Sc., Ph.D.")
    set_run_font(r, bold=True)

    core = doc.core_properties
    core.title = "Laporan Kemajuan Riset TISE 3.0 Tahun 2026"
    core.subject = "Program Riset P2MI STEI - Riset Guru Besar"
    core.author = "Prof. Ir. Armein Z. R. Langi, M.Sc., Ph.D."
    core.keywords = "TISE 3.0, laporan kemajuan, PPMI, publikasi, penggunaan dana"
    doc.save(DOCX_PATH)


class ProgressPdf(BaseDocTemplate):
    def __init__(self, filename, **kwargs):
        super().__init__(filename, **kwargs)
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="normal")
        self.addPageTemplates(PageTemplate(id="main", frames=frame, onPage=self.draw_header_footer))

    def draw_header_footer(self, canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#D7DBE2"))
        canvas.setLineWidth(0.5)
        canvas.line(doc.leftMargin, letter[1] - 0.62 * inch, letter[0] - doc.rightMargin, letter[1] - 0.62 * inch)
        canvas.setFont("Helvetica-Bold", 7.5)
        canvas.setFillColor(colors.HexColor("#5E6B75"))
        canvas.drawString(doc.leftMargin, letter[1] - 0.50 * inch, "LAPORAN KEMAJUAN RISET PPMI - RISET GURU BESAR 2026")
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(letter[0] - doc.rightMargin, 0.45 * inch, f"Halaman {doc.page}")
        canvas.restoreState()


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("TitleX", parent=base["Title"], fontName="Helvetica-Bold", fontSize=21, leading=24, textColor=colors.HexColor("#111111"), alignment=TA_LEFT, spaceAfter=5),
        "subtitle": ParagraphStyle("SubtitleX", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#1F4D78"), spaceAfter=15),
        "meta": ParagraphStyle("Meta", parent=base["Normal"], fontName="Helvetica", fontSize=9.4, leading=12, spaceAfter=2),
        "h1": ParagraphStyle("H1X", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=18, textColor=colors.HexColor("#2E74B5"), spaceBefore=13, spaceAfter=7, keepWithNext=True),
        "h2": ParagraphStyle("H2X", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=14, textColor=colors.HexColor("#2E74B5"), spaceBefore=9, spaceAfter=5, keepWithNext=True),
        "body": ParagraphStyle("BodyX", parent=base["BodyText"], fontName="Helvetica", fontSize=9.7, leading=12.2, textColor=colors.HexColor("#111111"), alignment=TA_JUSTIFY, spaceAfter=6),
        "small": ParagraphStyle("SmallX", parent=base["BodyText"], fontName="Helvetica", fontSize=7.6, leading=9.2, textColor=colors.HexColor("#111111")),
        "small_center": ParagraphStyle("SmallCenter", parent=base["BodyText"], fontName="Helvetica", fontSize=7.6, leading=9.2, textColor=colors.HexColor("#111111"), alignment=TA_CENTER),
        "th": ParagraphStyle("TH", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=7.7, leading=9, textColor=colors.HexColor("#1F4D78"), alignment=TA_CENTER),
        "note": ParagraphStyle("Note", parent=base["BodyText"], fontName="Helvetica-Oblique", fontSize=8.5, leading=10.5, textColor=colors.HexColor("#4F5B65"), spaceAfter=6),
        "callout_head": ParagraphStyle("CalloutHead", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=10, leading=12, textColor=colors.HexColor("#1F4D78"), spaceAfter=3),
        "callout": ParagraphStyle("Callout", parent=base["BodyText"], fontName="Helvetica", fontSize=9.2, leading=11.4, textColor=colors.HexColor("#111111")),
    }


def P(text: str, style):
    return Paragraph(text.replace("&", "&amp;"), style)


def pdf_table(headers, rows, widths, styles, font_size=7.6, header_fill="#F2F4F7"):
    header_cells = [P(h, styles["th"]) for h in headers]
    body_style = ParagraphStyle("tbl", parent=styles["small"], fontSize=font_size, leading=font_size + 1.5)
    center_style = ParagraphStyle("tblc", parent=body_style, alignment=TA_CENTER)
    data = [header_cells]
    for row in rows:
        formatted = []
        for idx, value in enumerate(row):
            st = center_style if idx == 0 and len(headers) > 3 else body_style
            formatted.append(P(str(value), st))
        data.append(formatted)
    table = Table(data, colWidths=[w * inch for w in widths], repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(header_fill)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B7C0C8")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def build_pdf() -> None:
    s = pdf_styles()
    doc = ProgressPdf(str(PDF_PATH), pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=0.82 * inch, bottomMargin=0.75 * inch, title="Laporan Kemajuan Riset TISE 3.0 Tahun 2026", author="Prof. Ir. Armein Z. R. Langi")
    story = []
    story.append(Spacer(1, 0.20 * inch))
    story.append(P("LAPORAN KEMAJUAN PROYEK RISET", s["title"]))
    story.append(P("Meta-Architecture of Triune-Intelligence Smart Engineering (TISE 3.0)", s["subtitle"]))
    for label, value in (
        ("Skema", "Program Riset P2MI STEI - Riset Guru Besar Tahun Anggaran 2026"),
        ("Kode file", "STEI.PPMI-1-43-2026"),
        ("Ketua peneliti", "Prof. Ir. Armein Z. R. Langi, M.Sc., Ph.D."),
        ("Unit", "Kelompok Keahlian Teknologi Informasi, STEI ITB"),
        ("Periode", "1 April - 31 Desember 2026"),
        ("Posisi laporan", "19 Agustus 2026; transaksi keuangan tersedia sampai 28 Juli 2026"),
    ):
        story.append(P(f"<b>{label}:</b> {value}", s["meta"]))
    story.append(Spacer(1, 8))
    callout = Table([[P("STATUS KESELURUHAN: BERJALAN SESUAI TAHAPAN, DENGAN RISIKO TERKENDALI", s["callout_head"]),], [P("Arsitektur, rancangan purwarupa, dan instrumen eksperimen telah terdokumentasi. Prioritas berikutnya adalah menghasilkan bukti simulasi final, mengisi bagian hasil, memastikan keterlacakan artefak, dan menuntaskan submit artikel Q1.", s["callout"]),]], colWidths=[6.5 * inch], hAlign="LEFT")
    callout.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF7E0")), ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#D7B75D")), ("LEFTPADDING", (0, 0), (-1, -1), 9), ("RIGHTPADDING", (0, 0), (-1, -1), 9), ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
    story.append(callout)

    story.append(P("1. Ringkasan Eksekutif", s["h1"]))
    story.append(P("Riset TISE 3.0 bertujuan membangun meta-arsitektur berbasis prompt yang mengubah niat pemangku kepentingan menjadi spesifikasi SysML v2 dan artefak TISE 2.0 melalui orkestrasi Meta-Architect, Generative Engineer, dan Critic. Berdasarkan dokumen yang tersedia, pekerjaan telah bergerak dari formulasi konsep menuju kesiapan validasi Digital Twin. [S1][S3]", s["body"]))
    metrics = [
        ("Kemajuan teknis", "Arsitektur dan desain eksperimen tersedia; hasil empiris final belum tersedia."),
        ("Luaran ilmiah", "Pipeline: 7 published, 1 under review, 2 draft; terdapat pula manuskrip SVCC 74 halaman yang relevan langsung."),
        ("Keuangan", "Realisasi Rp81.480.000 dari Rp150.000.000; serapan 54% menurut LPD (54,32% dari perhitungan)."),
        ("Target kontraktual", "Satu publikasi Q1; rencana submit 30 November 2026."),
    ]
    story.append(pdf_table(["Aspek", "Posisi Kemajuan"], metrics, [1.55, 4.95], s, font_size=8.6, header_fill="#E8EEF5"))

    story.append(P("2. Sasaran dan Posisi Terhadap Jadwal", s["h1"]))
    story.append(P("Proposal menetapkan tiga tujuan: memformulasikan TISE 3.0, membangun dan mensimulasikan purwarupa prompt-to-artifact, serta mempublikasikan hasil pada jurnal internasional Q1. Jadwal April-Desember menempatkan Agustus-September pada tahap Inner Right W-Model, yaitu persiapan validasi PICOC dan Digital Twin. [S1]", s["body"]))
    story.append(pdf_table(["Paket Kerja", "Target Proposal", "Bukti Kemajuan", "Status"], PROGRESS_ROWS, [1.25, 1.55, 2.75, 0.95], s, font_size=7.0))
    story.append(Spacer(1, 5))
    story.append(P("Catatan penilaian: status di atas diturunkan dari dokumen, bukan dari pemeriksaan repositori kode atau log eksperimen. Klaim implementasi yang disebut dalam SVCC perlu dilengkapi dengan versi kode, konfigurasi, seed, dan keluaran run yang dibekukan.", s["note"]))

    story.append(P("3. Capaian Teknis dan Ilmiah", s["h1"]))
    story.append(P("3.1 Meta-arsitektur TISE 3.0", s["h2"]))
    story.append(P("Manuskrip Systemic Value Co-Creation (SVCC) mendefinisikan TISE 3.0 sebagai lapisan design-time yang menghasilkan artefak sosio-teknis TISE 2.0. Transformasi disusun dalam lima tahap: prompt pemangku kepentingan, dekomposisi ASTF, formalisasi PSKVE/VOMR, spesifikasi SysML v2, dan instansiasi artefak. Peran agen dipisahkan agar keluaran antara, konflik, dan pemeriksaan dapat diaudit. [S3]", s["body"]))
    story.append(P("3.2 Purwarupa dan Digital Twin Mini-Lab", s["h2"]))
    story.append(P("SVCC mendeskripsikan paket Python berorientasi objek dengan Stakeholder, VOMR account, beberapa MIG reasoner, tiga coordinator, pasar Stock/Financial/Resource-Output, PSKVE Exchange Engine, Triune Intelligence, dan adapter penyedia LLM. Alur simulasi mencakup observasi, proposal MIG, koordinasi, pemeriksaan kendala, eksekusi, settlement, dan pencatatan EVI/resiliensi. Ini merupakan bukti kesiapan arsitektural dan implementatif, tetapi laporan ini belum memverifikasi berkas kode secara terpisah. [S3]", s["body"]))
    story.append(P("3.3 Validasi dan batas klaim", s["h2"]))
    story.append(P("Desain eksperimen membedakan efek penalaran stakeholder dari efek koordinasi dan membandingkan baseline, fixed-policy SOTA, serta adaptive TISE. Pertanyaan validasi, skenario normal/gangguan, metrik, dan format tabel hasil telah tersedia. Namun, bagian Results masih berupa kerangka yang akan diisi setelah Mini-Lab dijalankan; karena itu belum layak menyatakan peningkatan performa, resiliensi, atau efisiensi kompresi nilai sebagai temuan empiris. [S3]", s["body"]))

    story.append(P("4. Luaran Publikasi", s["h1"]))
    story.append(P("Dokumen Publication Pipeline 2026 mencatat sepuluh luaran: tujuh published, satu under review, dan dua draft. Dari tujuh yang published, dua ditandai Q1, dua Q2/Q3, dua Q3, dan satu prosiding non-Q. [S2]", s["body"]))
    story.append(pdf_table(["Status", "Jumlah", "Rincian"], [("Published", "7", "2 Q1; 2 Q2/Q3; 2 Q3; 1 prosiding non-Q"), ("Under review", "1", "Q1 - IEEE Transactions on Education"), ("Draft", "2", "Keduanya diarahkan ke jurnal Q1 IEEE")], [1.35, 0.75, 4.4], s, font_size=8.3, header_fill="#E8EEF5"))
    story.append(Spacer(1, 5))
    story.append(P("Kesesuaian dengan target proposal: target kontraktual adalah satu artikel bertopik TISE 3.0 untuk jurnal Q1 dengan rencana submit 30 November 2026. Tidak ada entri pipeline yang sama persis dengan judul target. Manuskrip SVCC memiliki keterkaitan langsung dan isi yang memadai sebagai basis artikel, tetapi status jurnal/submit belum dinyatakan dan bagian hasil masih menunggu data. Oleh sebab itu target publikasi proyek dinilai masih dalam proses, meskipun pipeline publikasi peneliti secara umum kuat. [S1][S2][S3]", s["body"]))

    story.append(PageBreak())
    story.append(P("4.1 Daftar Publication Pipeline 2026", s["h2"]))
    pub_rows = [(str(no), title, journal, quartile, status) for no, title, journal, quartile, status in PUBLICATIONS]
    story.append(pdf_table(["No", "Judul", "Target jurnal/forum", "Kuartil", "Status"], pub_rows, [0.35, 3.05, 1.45, 0.65, 1.0], s, font_size=6.7))
    story.append(Spacer(1, 5))
    story.append(P("Catatan: daftar ini direproduksi dari pipeline yang disediakan. Laporan tidak mengasumsikan bahwa seluruh entri merupakan luaran langsung pendanaan PPMI TISE 3.0 tanpa bukti acknowledgement atau pemetaan biaya. [S2]", s["note"]))

    story.append(P("5. Penggunaan Dana", s["h1"]))
    story.append(P("LPD mencatat penerimaan Rp150.000.000 dan 13 transaksi pengeluaran dengan total Rp81.480.000. Saldo Rp68.520.000 dan presentase serapan dilaporkan 54%. Perhitungan terhadap pagu menghasilkan 54,32%, sehingga perbedaan hanya akibat pembulatan. [S4]", s["body"]))
    story.append(pdf_table(["Indikator", "Nilai"], [("Pagu/penerimaan", rupiah(150_000_000)), ("Realisasi", rupiah(81_480_000)), ("Saldo", rupiah(68_520_000)), ("Serapan", "54% (LPD); 54,32% (perhitungan)"), ("Cakupan transaksi", "28 April - 28 Juli 2026; laporan dinyatakan Agustus 2026")], [2.0, 4.5], s, font_size=8.4, header_fill="#E8EEF5"))
    story.append(Spacer(1, 7))
    finance_table = [(c, rupiah(p) if p else "Tidak terpisah", rupiah(a), x, n) for c, p, a, x, n in FINANCE_ROWS]
    finance_table.append(("Total", rupiah(150_000_000), rupiah(81_480_000), "54,32%", "Saldo keseluruhan " + rupiah(68_520_000) + "."))
    story.append(pdf_table(["Kategori", "RAB Proposal", "Realisasi LPD", "Serapan", "Keterangan"], finance_table, [1.05, 1.2, 1.2, 0.75, 2.3], s, font_size=7.1))
    story.append(Spacer(1, 5))
    story.append(P("Catatan rekonsiliasi: perbandingan kategori bersifat indikatif. RAB proposal menggabungkan honor mahasiswa, perjalanan, AI/hosting, internet/cloud, serta registrasi/publikasi di kelompok jasa; LPD menggunakan klasifikasi pegawai, barang, jasa, dan modal. Belanja modal Rp10.400.000 perlu dipetakan secara administratif ke pos RAB yang disetujui sebelum laporan akhir. [S1][S4]", s["note"]))

    story.append(P("6. Isu, Risiko, dan Tindakan Korektif", s["h1"]))
    risks = [("Bukti empiris belum final", "Tinggi", "Bekukan protokol; jalankan repeated matched runs; laporkan ketidakpastian, sensitivitas, ablation, dan kasus gagal."), ("Target Q1 belum terdokumentasi sebagai submitted", "Tinggi", "Tetapkan artikel/jurnal; selesaikan Results/Discussion; periksa format dan acknowledgement; submit sebelum 30 November."), ("Keterlacakan prompt-SysML-artefak", "Sedang", "Sediakan trace matrix, versi model/kode, konfigurasi provider, seed, log, dan artefak keluaran."), ("Klasifikasi belanja modal", "Sedang", "Rekonsiliasi dengan RAB dan bagian keuangan; dokumentasikan justifikasi/persetujuan."), ("Batas klaim simulasi", "Sedang", "Pisahkan bukti simulasi dari kondisi kampus nyata; jelaskan asumsi, kalibrasi, dan validasi eksternal.")]
    story.append(pdf_table(["Isu/Risiko", "Prioritas", "Tindakan"], risks, [1.65, 0.8, 4.05], s, font_size=7.4))

    story.append(P("7. Rencana Kerja Sampai Akhir Tahun", s["h1"]))
    story.append(pdf_table(["Periode", "Keluaran yang Diharapkan"], WORKPLAN_ROWS, [1.45, 5.05], s, font_size=8.0, header_fill="#E8EEF5"))

    story.append(P("8. Kesimpulan", s["h1"]))
    story.append(P("Secara keseluruhan, riset berada pada posisi yang konsisten dengan peralihan menuju tahap validasi Agustus-September. Fondasi konseptual, rancangan meta-arsitektur, struktur purwarupa, dan protokol Digital Twin telah terdokumentasi dengan baik. Kesenjangan utama adalah belum tersedianya hasil eksperimen final dan belum adanya bukti bahwa artikel TISE 3.0 telah masuk proses submit. Dengan serapan anggaran 54% dan sisa periode sampai Desember, sasaran masih dapat dicapai apabila eksekusi eksperimen, penyusunan artikel, serta rekonsiliasi anggaran diprioritaskan segera.", s["body"]))

    story.append(PageBreak())
    story.append(P("Lampiran: Sumber Data dan Catatan Verifikasi", s["h1"]))
    sources = [("[S1]", "proposal5signed.pdf - proposal 18 halaman; identitas, tujuan, target publikasi, jadwal, dan RAB."), ("[S2]", "bahan/pipeline2026.pdf - Publication Pipeline 2026, 2 halaman."), ("[S3]", "bahan/SVCC.pdf - Systemic Value Co-Creation, 74 halaman, bertanggal 19 Agustus 2026."), ("[S4]", "LPD Armein Z.R Langi.pdf - Laporan Penggunaan Dana, 2 halaman, periode 1 April-31 Desember 2026.")]
    story.append(pdf_table(["Kode", "Dokumen"], sources, [0.65, 5.85], s, font_size=8.0))
    story.append(Spacer(1, 5))
    story.append(P("Metode: angka dan status ditranskripsi dari dokumen yang diberikan dan diperiksa terhadap tampilan halaman PDF. Tidak dilakukan verifikasi eksternal atas indeks jurnal, tautan bukti transfer, atau keberadaan kode sumber. Istilah published, under review, dan draft mengikuti pipeline. Status SVCC dinilai sebagai manuskrip pra-submisi karena tidak mencantumkan jurnal/status submit dan bagian Results menyatakan data belum tersedia.", s["note"]))
    story.append(Spacer(1, 14))
    sig = Table([["", P("Bandung, 19 Agustus 2026<br/>Ketua Peneliti,<br/><br/><br/><b>Prof. Ir. Armein Z. R. Langi, M.Sc., Ph.D.</b>", s["meta"])]], colWidths=[3.65 * inch, 2.85 * inch])
    sig.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("ALIGN", (1, 0), (1, 0), "LEFT")]))
    story.append(sig)
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
